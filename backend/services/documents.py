import os
import re
import json
import uuid
import time
import requests
from datetime import datetime, timezone
from typing import Generator

from fastapi import HTTPException
from sqlalchemy.orm import Session
from dotenv import load_dotenv

from db.models import Document
from db.mongo import summaries_collection
from schemas.documents_schemas import DocumentCreateResponse

load_dotenv()

ALLOWED_EXTENSIONS = {"pdf", "docx", "txt"}
MAX_FILE_SIZE = 21 * 1024 * 1024  # 21MB to handle files slightly over 20MB

UPLOAD_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)


# --- Document CRUD ---

def get_documents(db: Session):
    return db.query(Document).all()


def get_document_by_id(db: Session, doc_id: str):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


async def upload_document_service(file, title, fileType, db: Session):
    file_extension = file.filename.split(".")[-1].lower()

    if file_extension not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=422, detail="File extension not allowed")

    content = await file.read()

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=422,
            detail=f"File too large. Maximum size is 20MB. Your file is {len(content) // 1024}KB."
        )

    doc_id = str(uuid.uuid4())
    safe_filename = re.sub(r"[^a-zA-Z0-9_.-]", "_", file.filename)
    file_name = f"{doc_id}_{safe_filename}"
    full_path = os.path.join(UPLOAD_DIR, file_name)

    with open(full_path, "wb") as f_disk:
        f_disk.write(content)

    new_doc = Document(
        id=doc_id,
        title=title,
        fileType=file_extension,
        filePath=full_path,
        status="pending",
        fileSize=len(content),
    )
    db.add(new_doc)
    db.commit()
    db.refresh(new_doc)

    return DocumentCreateResponse(id=doc_id, message="Document uploaded successfully")


def delete_document_service(db: Session, doc_id: str):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.filePath and os.path.exists(doc.filePath):
        try:
            os.remove(doc.filePath)
        except Exception:
            raise HTTPException(status_code=500, detail="Could not delete file")

    db.delete(doc)
    db.commit()
    return {"message": "Document deleted successfully"}


# --- Text extraction ---

def extract_text_from_file(file_path: str, file_type: str) -> str:
    file_type = file_type.lower()

    if file_type == "txt":
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif file_type == "pdf":
        import fitz
        parts = []
        with fitz.open(file_path) as doc:
            for page in doc:
                parts.append(page.get_text())
        return "\n".join(parts)
    elif file_type == "docx":
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_type}")


# --- Chunking ---

def chunk_text(text: str, chunk_size: int = 4000, overlap: int = 400) -> list[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunk = text[start:end]
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


# --- Groq API call (non-streaming, used for map step) ---

def _call_groq(prompt: str, max_tokens: int = 400) -> str:
    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        return ""

    payload = {
        "model": "llama-3.1-8b-instant",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": 0.3,
        "max_tokens": max_tokens,
        "stream": False,
    }
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
    }

    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=payload,
            timeout=60,
        )

        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"].strip()
        elif response.status_code == 429:
            print("[Groq] Rate limited on chunk call, waiting 15s...")
            time.sleep(15)
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=60,
            )
            if response.status_code == 200:
                return response.json()["choices"][0]["message"]["content"].strip()
        else:
            print(f"[Groq chunk ERROR] {response.status_code}: {response.text[:200]}")

    except Exception as e:
        print(f"[Groq chunk exception] {e}")

    return ""


# --- Map step: summarize one chunk ---

def _summarize_chunk(chunk: str, chunk_index: int, total_chunks: int) -> str:
    prompt = (
        f"This is part {chunk_index + 1} of {total_chunks} from a study document. "
        "Extract the most important facts, definitions, concepts, dates, names, and key points "
        "from this section. Write them as a compact list of notes. "
        "Be specific — include actual values, names, and details. "
        "Do NOT include generic statements. Write in the SAME LANGUAGE as the source.\n\n"
        f"Section text:\n{chunk}\n\nKey points from this section:"
    )
    return _call_groq(prompt, max_tokens=350)


# --- Map-reduce summarization with streaming reduce step ---

def stream_summary_tokens(text: str, summary_type: str) -> Generator[str, None, None]:
    if not text.strip():
        raise HTTPException(status_code=400, detail="Document appears to be empty")

    api_key = os.getenv("GROQ_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="GROQ_API_KEY not set in .env")

    # Warn for very large documents
    text_length = len(text)
    if text_length > 450000:
        yield (
            f"⚠️ Note: This document is large ({text_length // 1000}K characters). "
            f"The summary will cover key sections but may not capture every detail.\n\n"
        )

    # --- MAP STEP ---
    chunks = chunk_text(text, chunk_size=4000, overlap=400)
    print(f"[MapReduce] {len(chunks)} chunks for {summary_type} summary")

    max_chunks = 5 if summary_type == "concise" else 6
    if len(chunks) > max_chunks:
        step = len(chunks) / max_chunks
        chunks = [chunks[int(i * step)] for i in range(max_chunks)]

    map_start = time.time()
    chunk_summaries = []

    for i, chunk in enumerate(chunks):
        print(f"[MapReduce] Summarizing chunk {i+1}/{len(chunks)}...")
        summary = _summarize_chunk(chunk, i, len(chunks))
        if summary:
            chunk_summaries.append(summary)
        if i < len(chunks) - 1:
            time.sleep(2)

    map_elapsed = time.time() - map_start

    if not chunk_summaries:
        yield fallback_summary(text, summary_type)
        return

    trimmed = [s[:200] for s in chunk_summaries]
    combined = "\n\n---\n\n".join(trimmed)

    # Wait only as long as needed for the TPM window to reset
    tpm_window = 75 if summary_type == "detailed" else 62
    wait_needed = max(0, tpm_window - map_elapsed)
    if wait_needed > 0:
        print(f"[MapReduce] Map took {map_elapsed:.0f}s, waiting {wait_needed:.0f}s for TPM reset...")
        time.sleep(wait_needed)
    else:
        print(f"[MapReduce] Map took {map_elapsed:.0f}s, no wait needed.")

    # --- REDUCE STEP (streaming) ---
    if summary_type == "concise":
        reduce_prompt = (
            "You are a university student writing study notes before an exam. "
            "Below are key points extracted from different sections of a document. "
            "Combine them into 10-15 clear bullet points covering the most important concepts.\n\n"
            "Rules:\n"
            "- Each bullet is 1-2 sentences with enough context to understand on its own\n"
            "- Include specific names, dates, numbers, and terminology\n"
            "- Write in the SAME LANGUAGE as the notes below\n"
            "- Plain text only, absolutely no markdown formatting (no **, ##, *)\n"
            "- Start directly with the first bullet point\n"
            "- Do NOT copy bullet points verbatim — synthesize and rewrite\n"
            "- Do NOT add any introduction or closing remark\n\n"
            f"Extracted notes from each section:\n{combined}\n\nFinal study notes:"
        )
    else:
        reduce_prompt = (
            "You are a university student writing a comprehensive study guide before an exam. "
            "Below are key points extracted from different sections of a document. "
            "Combine them into a thorough study guide organized by topic.\n\n"
            "Rules:\n"
            "- Organize into 4-6 sections with clear topic headings\n"
            "- Under each heading write 2-3 detailed paragraphs explaining concepts as if teaching a classmate\n"
            "- Include ALL key definitions, names, dates, examples, and important details\n"
            "- Keep total length under 650 words so you finish properly — do not cut off mid-sentence\n"
            "- Write a proper concluding paragraph at the end\n"
            "- Write in the SAME LANGUAGE as the notes below\n"
            "- Plain text only, absolutely no bullet points, no numbered lists, "
            "no dashes as list items, no 'Key Points' sections\n"
            "- If the notes contain bullet points, rewrite them as flowing paragraphs\n"
            "- Start directly with the first section heading\n"
            "- Do NOT add any introduction or closing remark before the first heading\n"
            "- Avoid repeating the same phrases — vary your language\n\n"
            f"Extracted notes from each section:\n{combined}\n\nFinal study guide:"
        )

    print("[MapReduce] Streaming reduce step...")
    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json={
                "model": "llama-3.3-70b-versatile" if summary_type == "concise" else "llama-3.1-8b-instant",
                "messages": [{"role": "user", "content": reduce_prompt}],
                "temperature": 0.4,
                "max_tokens": 800 if summary_type == "concise" else 1800,
                "stream": True,
            },
            stream=True,
            timeout=120,
        )

        print(f"[MapReduce reduce] Status: {response.status_code}")
        if response.status_code == 429:
            print("[MapReduce] Reduce step rate limited, waiting 20s...")
            time.sleep(20)
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json={
                    "model": "llama-3.1-8b-instant",
                    "messages": [{"role": "user", "content": reduce_prompt}],
                    "temperature": 0.4,
                    "max_tokens": 1800,
                    "stream": True,
                },
                stream=True,
                timeout=120,
            )
        if response.status_code != 200:
            print(f"[MapReduce reduce ERROR] {response.text[:300]}")
            yield fallback_summary(text, summary_type)
            return

        for line in response.iter_lines():
            if not line:
                continue
            decoded = line.decode("utf-8") if isinstance(line, bytes) else line
            if not decoded.startswith("data:"):
                continue
            raw = decoded[5:].strip()
            if raw == "[DONE]":
                break
            try:
                chunk_data = json.loads(raw)
                token = chunk_data["choices"][0]["delta"].get("content", "")
                if token:
                    yield token
            except (json.JSONDecodeError, IndexError, KeyError):
                continue

    except Exception as e:
        print(f"[MapReduce reduce exception] {e}")
        yield fallback_summary(text, summary_type)


# --- Clean AI output ---

def _clean_ai_output(text: str) -> str:
    text = text.replace("**", "").replace("##", "").replace("###", "")
    text = re.sub(r"(?<!\w)\*([^*]+)\*(?!\w)", r"\1", text)

    preamble_patterns = [
        r"^(Sure[!,.]?\s*)",
        r"^(Of course[!,.]?\s*)",
        r"^(Absolutely[!,.]?\s*)",
        r"^(Here (?:is|are) .*?[:.]\s*\n*)",
        r"^(Below (?:is|are) .*?[:.]\s*\n*)",
        r"^(The following .*?[:.]\s*\n*)",
        r"^(These are .*?[:.]\s*\n*)",
        r"^(This is .*?summary.*?[:.]\s*\n*)",
        r"^(I'?ve (?:compiled|created|written|prepared|summarized) .*?[:.]\s*\n*)",
    ]
    for pattern in preamble_patterns:
        text = re.sub(pattern, "", text, count=1, flags=re.IGNORECASE)

    signoff_patterns = [
        r"\n+(?:I hope this helps|Let me know if|Feel free to|"
        r"If you have any|Don't hesitate|Happy studying|"
        r"Good luck|Best of luck).*$",
    ]
    for pattern in signoff_patterns:
        text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.DOTALL)

    return text.strip()


def fallback_summary(text: str, summary_type: str) -> str:
    sentences = [
        s.strip()
        for s in text.replace("\n", " ").split(".")
        if len(s.strip()) > 20
    ]
    count = 5 if summary_type == "concise" else 12
    return ". ".join(sentences[:count]) + "." if sentences else "Could not generate summary."


# --- Summary service functions ---

async def summarize_document_service(db: Session, doc_id: str, summary_type: str):
    doc = db.query(Document).filter(Document.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    text = extract_text_from_file(doc.filePath, doc.fileType)
    raw = "".join(stream_summary_tokens(text, summary_type))
    raw = re.sub(r"^⚠️[^\n]+\n\n?", "", raw)
    summary = _clean_ai_output(raw)

    await summaries_collection.update_one(
        {"documentId": doc_id},
        {
            "$set": {
                "documentId": doc_id,
                f"summaries.{summary_type}": {
                    "summary": summary,
                    "generatedAt": datetime.utcnow(),
                },
            }
        },
        upsert=True,
    )

    doc.status = "summarized"
    db.commit()

    return {
        "message": "Summarization complete",
        "documentId": doc_id,
    }


async def get_summary_service(db, doc_id: str, summary_type: str = "concise"):
    result = await summaries_collection.find_one({"documentId": doc_id})

    if not result or "summaries" not in result or summary_type not in result["summaries"]:
        raise HTTPException(
            status_code=404,
            detail=f"No {summary_type} summary found. Generate one first."
        )

    type_data = result["summaries"][summary_type]
    return {
        "documentId": result["documentId"],
        "title": get_document_by_id(db, doc_id).title,
        "summary": type_data["summary"],
        "summaryType": summary_type,
        "generatedAt": type_data["generatedAt"],
    }


async def regenerate_summary_service(db, doc_id: str, summary_type: str):
    return await summarize_document_service(db, doc_id, summary_type)